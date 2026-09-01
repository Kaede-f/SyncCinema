"""Build the SyncCinema engineering journal DOCX.

The DOCX is committed together with this source file:
- the Python source keeps document changes reviewable in Git;
- the DOCX gives the project owner a convenient learning/interview artifact.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).with_name("SyncCinema_Engineering_Journal.docx")

FONT_BODY = "Microsoft YaHei"
FONT_CODE = "Consolas"
COLOR_BLUE = RGBColor(46, 116, 181)
COLOR_DARK_BLUE = RGBColor(31, 77, 120)
COLOR_INK = RGBColor(30, 38, 48)
COLOR_MUTED = RGBColor(91, 101, 113)
COLOR_LIGHT_BLUE = "E8EEF5"
COLOR_LIGHT_GRAY = "F2F4F7"
COLOR_CALLOUT = "F4F6F9"
COLOR_BORDER = "C8D2DE"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    name=FONT_BODY,
    size=10.5,
    color=COLOR_INK,
    bold=None,
    italic=None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLOR_BORDER, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_geometry(table, column_widths_dxa):
    """Apply fixed OOXML geometry so Word/LibreOffice render identically."""
    if sum(column_widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError("Table column widths must total 9360 DXA")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in column_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = column_widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

    set_table_borders(table)


def mark_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = COLOR_INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, COLOR_BLUE, 18, 10),
        "Heading 2": (13, COLOR_BLUE, 14, 7),
        "Heading 3": (12, COLOR_DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header_paragraph, after=0, line_spacing=1.0)
    run = header_paragraph.add_run("SyncCinema  |  工程演进与面试复盘")
    set_run_font(run, size=8.5, color=COLOR_MUTED, bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_paragraph, after=0, line_spacing=1.0)
    run = footer_paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=COLOR_MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)
    run = footer_paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=COLOR_MUTED)


def add_title_block(doc):
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=12)

    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, after=4, line_spacing=1.0)
    run = kicker.add_run("ENGINEERING JOURNAL")
    set_run_font(run, size=9.5, color=COLOR_BLUE, bold=True)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, after=6, line_spacing=1.0)
    run = title.add_run("SyncCinema 工程演进与面试复盘")
    set_run_font(run, size=25, color=COLOR_DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=18, line_spacing=1.15)
    run = subtitle.add_run("从命令行同步观影 MVP 到可测量、可校正、可产品化的多人观影系统")
    set_run_font(run, size=12.5, color=COLOR_MUTED)

    metadata = [
        ("当前开发分支", "feature/qt-client-ui"),
        ("当前已发布提交", "2f4f27a - add read-only sync correction advice"),
        ("本次阶段", "Qt 桌面客户端、共享会话层与播放器画面嵌入"),
        ("维护日期", "2026-07-29"),
        ("安全约束", "服务器口令、私钥和其他凭据不得进入源码、日志、文档或 Git 历史"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    for index, header in enumerate(("项目元数据", "当前值")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for label, value in metadata:
        label_cell, value_cell = table.add_row().cells
        set_cell_shading(label_cell, COLOR_LIGHT_BLUE)
        for cell, text, bold in (
            (label_cell, label, True),
            (value_cell, value, False),
        ):
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5, bold=bold)

    note = doc.add_paragraph()
    set_paragraph_spacing(note, before=14, after=4, line_spacing=1.25)
    run = note.add_run("文档定位：")
    set_run_font(run, bold=True, color=COLOR_DARK_BLUE)
    run = note.add_run(
        "每个工程阶段都记录问题背景、设计判断、核心代码、验证证据、已知局限和面试表达。"
        "后续即使暂时不参与编码，也能据此重新进入项目。"
    )
    set_run_font(run)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_run_font(prefix_run, bold=True, color=COLOR_DARK_BLUE)
        text = text[len(bold_prefix):]
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=6, after=8, line_spacing=1.25)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), COLOR_CALLOUT)
    p_pr.append(shading)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "160")
    ind.set(qn("w:right"), "160")
    p_pr.append(ind)
    run = paragraph.add_run(f"{label}：")
    set_run_font(run, bold=True, color=COLOR_DARK_BLUE)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=4, after=8, line_spacing=1.1)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F7F8FA")
    p_pr.append(shading)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "180")
    p_pr.append(ind)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, name=FONT_CODE, size=8.5, color=COLOR_INK)
        if index < len(lines) - 1:
            run.add_break()
    return paragraph


def add_commit_history_table(doc):
    rows = [
        ("5b76dc3", "Initial SyncCinema libVLC MVP", "真实本地视频可被 play/pause/seek 控制"),
        ("991a4f4", "multi-client broadcast architecture", "server 协调多个 client 并广播控制命令"),
        ("3483a28", "cross-platform Linux server build", "Linux 云端可单独构建轻量 server"),
        ("d0d91e3", "stabilize cross-platform builds", "统一源码编码与 Windows/Linux 构建行为"),
        ("90e5e3e", "HTTP media source", "client 可播放 Nginx 提供的云端视频"),
        ("ab2adc8", "client startup timings", "量化播放器初始化、媒体打开和 TCP 连接耗时"),
        ("c7caec9", "progress drift analysis", "上报真实播放器进度并比较 Room 理论进度"),
        ("02f6aba", "heartbeat RTT measurements", "以 PING/PONG 可靠测量每个 client 的 RTT"),
        ("2c1bf4f", "pairwise playback skew analysis", "把两端上报投影到同一时刻，直接测量客户端间偏差"),
        ("4944cf8", "control epochs and robust skew windows", "隔离控制周期并以中位数/P95 过滤单点抖动"),
        ("92300ce", "late-joining room snapshots", "新 client 加入时按毫秒对齐当前房间状态与控制版本"),
        ("2f4f27a", "read-only sync correction advice", "以安全门和可解释原因生成校正建议，但不改变播放行为"),
    ]

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1450, 3300, 4610])
    headers = ("提交", "阶段", "可验证结果")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for commit, stage, result in rows:
        cells = table.add_row().cells
        for index, text in enumerate((commit, stage, result)):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                name=FONT_CODE if index == 0 else FONT_BODY,
                size=8.7 if index == 0 else 9,
            )


def add_changed_files_table(doc):
    rows = [
        (
            "SyncMetrics.h",
            "为每个 client 保存最近一次位置、状态和 server 接收时刻。",
            "让下一份上报到达时能找到可比较的另一端快照。",
        ),
        (
            "SyncMetrics.cpp",
            "新增统一时刻投影、陈旧样本过滤、状态一致性检查和 pair_progress 日志。",
            "直接得到 client A 与 client B 的相对播放偏差。",
        ),
        (
            "build_engineering_journal.py",
            "用可审查源码生成本 DOCX，固定版式和内容结构。",
            "后续每次提交都能持续维护学习与面试材料。",
        ),
    ]

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 3580, 3580])
    headers = ("文件", "核心修改", "工程价值")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for file_name, change, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((file_name, change, value)):
            paragraph = cells[index].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                name=FONT_CODE if index == 0 else FONT_BODY,
                size=8.6 if index == 0 else 9,
            )


def add_epoch_changed_files_table(doc):
    rows = [
        (
            "SyncServer.cpp",
            "新增 controlCommandMutex，使控制命令、Room 更新、广播、epoch 切换和 REPORT 快照严格有序。",
            "避免并发 client 产生“新状态配旧统计周期”的竞态。",
        ),
        (
            "SyncMetrics.h",
            "新增 PairWindowStats、控制 epoch 状态及 beginControlEpoch()。",
            "把网络 RTT 与播放控制周期分开管理。",
        ),
        (
            "SyncMetrics.cpp",
            "控制后保留 RTT、重置偏差；加入 2 秒稳定期、12 样本窗口、中位数、P95 和连续严重偏差计数。",
            "让后续纠偏依据稳定趋势，而不是单次量化抖动。",
        ),
        (
            "Engineering Journal",
            "补录公网验证证据、算法选择、回归结果和面试表达。",
            "保证代码推进与知识交接同步完成。",
        ),
    ]

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 3580, 3580])
    headers = ("文件/模块", "核心修改", "工程价值")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for module, change, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((module, change, value)):
            paragraph = cells[index].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                name=FONT_CODE if index == 0 else FONT_BODY,
                size=8.6 if index == 0 else 9,
            )


def add_snapshot_changed_files_table(doc):
    rows = [
        (
            "Protocol.h/.cpp",
            "新增 SNAPSHOT 消息、controlEpoch、严格解析与控制消息分类。",
            "把晚加入状态变成可测试、可演进的正式协议，而不是临时字符串。",
        ),
        (
            "Room.h/.cpp",
            "新增原子 RoomSnapshot；每条有效控制命令递增唯一 epoch。",
            "状态、位置和版本在同一把锁下读取，不会拼出撕裂快照。",
        ),
        (
            "SyncServer.cpp",
            "注册 client、读取快照和发送快照与控制命令共用顺序锁。",
            "新 client 不会在加入与初始化之间漏掉 PLAY、PAUSE 或 SEEK。",
        ),
        (
            "Client / PlayerController",
            "新增 5 秒握手、PING 兼容、缓冲保留、seekable 等待和毫秒级 seek。",
            "真实 libVLC 准备完成后再按快照精确定位，避免整秒取整误差。",
        ),
        (
            "CMake / core tests",
            "新增无第三方测试目标，覆盖协议、Room epoch 和伪造快照拒绝。",
            "Windows 与 Linux 都能快速锁住核心行为。",
        ),
    ]

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 3580, 3580])
    headers = ("文件/模块", "核心修改", "工程价值")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for module, change, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((module, change, value)):
            paragraph = cells[index].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                name=FONT_CODE if index == 0 else FONT_BODY,
                size=8.6 if index == 0 else 9,
            )


def add_correction_changed_files_table(doc):
    rows = [
        (
            "SyncCorrectionPolicy.h/.cpp",
            "新增纯策略输入、配置、安全门、动作与原因；策略不访问 socket、Room 或播放器。",
            "测量与决策解耦，可独立单元测试，也不会误触真实 seek。",
        ),
        (
            "SyncMetrics.h/.cpp",
            "从稳健窗口构造策略输入，计算方向一致率，并输出去重后的 correction_advice。",
            "把既有指标转成可解释的控制建议，同时保留完整观测证据。",
        ),
        (
            "SyncCinemaCoreTests",
            "覆盖稳定期、窗口、状态、RTT、阈值、方向、连续性、冷却及目标端选择。",
            "为未来闭环执行提供快速、跨平台的策略回归保护。",
        ),
        (
            "CMakeLists.txt",
            "让 Windows client、Linux server 和核心测试共同编译策略模块。",
            "保证同一份决策规则在所有目标中保持一致。",
        ),
        (
            "Engineering Journal",
            "记录策略边界、安全门、本地证据、已知限制和面试表达。",
            "让后续学习、评审与算法迭代具有连续上下文。",
        ),
    ]

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 3580, 3580])
    headers = ("文件/模块", "核心修改", "工程价值")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for module, change, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((module, change, value)):
            paragraph = cells[index].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                name=FONT_CODE if index == 0 else FONT_BODY,
                size=8.6 if index == 0 else 9,
            )


def add_qt_changed_files_table(doc):
    rows = [
        (
            "SyncClientSession.h/.cpp",
            "从原 Client.cpp 抽出连接、快照握手、广播接收、PING/PONG、进度上报和串行发送。",
            "CLI 与 Qt 共用同一套网络行为，后续协议升级不再维护两份实现。",
        ),
        (
            "QtClientWindow.h/.cpp",
            "新增媒体/服务器输入、连接状态、视频区域、播放控制、进度、音量和全屏界面。",
            "普通测试用户无需学习命令行，即可完成核心观影流程。",
        ),
        (
            "QtClientController.h/.cpp",
            "在后台线程初始化 libVLC 和连接 server，并用 queued invoke 把网络回调送回 UI 线程。",
            "避免冷启动和网络握手冻结界面，也禁止后台线程直接操作 QWidget。",
        ),
        (
            "PlayerController / LibVlcPlayer",
            "补充原生视频窗口、媒体总时长和音量接口，并用 libvlc_media_player_set_hwnd 嵌入画面。",
            "播放器能力继续隐藏在抽象接口后，网络会话不依赖 libVLC C API。",
        ),
        (
            "CMake / presets",
            "新增可选 SyncCinemaQt 目标、Qt5/Qt6 发现、windeployqt 与 VLC runtime 自动复制。",
            "Qt 默认关闭，Windows CLI 与 Linux server 保持原构建路径；发布目录可独立运行。",
        ),
        (
            "QtMsvcCompatibility.h",
            "仅在 Qt 5 + 新版 MSVC 下替换已经被 VS 2026 删除的 stdext 迭代器宏。",
            "兼容当前开发环境；迁移到官方 Qt 6 后自动退出构建。",
        ),
    ]

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 3580, 3580])
    headers = ("文件/模块", "核心修改", "工程价值")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLOR_LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=COLOR_DARK_BLUE)
    mark_repeat_table_header(table.rows[0])

    for module, change, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((module, change, value)):
            paragraph = cells[index].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                name=FONT_CODE if index == 0 else FONT_BODY,
                size=8.6 if index == 0 else 9,
            )


def build_document():
    doc = Document()
    doc.core_properties.title = "SyncCinema 工程演进与面试复盘"
    doc.core_properties.subject = "SyncCinema engineering journal"
    doc.core_properties.author = "SyncCinema Project"
    doc.core_properties.last_modified_by = "SyncCinema Project"
    doc.core_properties.keywords = "C++, TCP, libVLC, synchronization, metrics"
    configure_document(doc)
    add_title_block(doc)

    doc.add_heading("1. 当前项目全景", level=1)
    add_body(
        doc,
        "SyncCinema 当前已经不是最初的单连接 TCP demo，而是一个可以在公网运行的多人同步观影原型。"
        "视频内容通过 HTTP 分发，TCP server 只传输体积很小的控制与测量消息。"
    )
    add_callout(
        doc,
        "当前架构",
        "Windows client A/B 使用 libVLC 播放同一媒体 URL；Linux 云服务器上的 "
        "SyncCinemaServer 监听 9000 端口并维护 Room；Nginx 负责 80/443 端口的静态视频分发。",
    )
    add_code_block(
        doc,
        [
            "Windows Client A --\\",
            "                   +-- TCP :9000 --> Linux SyncCinemaServer / Room",
            "Windows Client B --/                         |",
            "                                                +-- metrics / broadcast",
            "",
            "Windows Client A/B ---- HTTP :80/443 ----> Nginx video files",
        ],
    )
    add_body(
        doc,
        "职责边界：Client 负责真实播放器和用户交互；Room 负责权威播放状态；"
        "SyncMetricsCollector 负责观测事实；Protocol 负责结构化消息与文本协议；"
        "Nginx 负责媒体内容，而不是 C++ server。",
    )

    doc.add_heading("2. 已完成的工程里程碑", level=1)
    add_body(
        doc,
        "下面的提交形成一条可解释的项目演进链。每一步都解决一个独立问题，"
        "因此既便于回滚，也方便面试时说明设计为什么逐步变化。"
    )
    add_commit_history_table(doc)

    doc.add_heading("3. 本次阶段：客户端间相对播放偏差", level=1)
    doc.add_heading("3.1 为什么原指标还不够", level=2)
    add_body(
        doc,
        "此前的 compensated_diff_ms 比较“某个 client 的 libVLC 进度”和“Room 的理论进度”。"
        "Room 在收到 PLAY 后立即按 steady_clock 前进，而 libVLC 可能先经历网络缓冲、解码和渲染准备。"
        "如果两台播放器都比 Room 晚 400 ms，但彼此只差 20 ms，旧指标仍会把两台都标成 watch。"
    )
    add_callout(
        doc,
        "核心判断",
        "用户真正感知的是两位观众看到的画面是否一致。因此自动同步算法的主要输入必须是 "
        "client-to-client 相对偏差，而 client-to-Room 偏差更适合诊断整体启动或播放器延迟。",
    )

    doc.add_heading("3.2 时间归一化算法", level=2)
    add_body(
        doc,
        "两台 client 每秒上报一次，消息不会同时到达 server。不能直接拿两条原始 position 相减，"
        "否则较早到达的样本天然更小。server 以自己的 steady_clock 为统一时间轴，"
        "把两份样本投影到同一个比较时刻。"
    )
    add_code_block(
        doc,
        [
            "Playing:",
            "projected_position = reported_position",
            "                   + estimated_one_way_delay",
            "                   + (compare_time - report_received_time)",
            "",
            "Paused / Stopped:",
            "projected_position = reported_position",
            "",
            "pair_diff_ms = projected_client_a - projected_client_b",
        ],
    )
    add_body(
        doc,
        "estimated_one_way_delay 暂时采用 min_rtt / 2。最小 RTT 更接近链路的基础传播延迟，"
        "比包含瞬时排队抖动的平均 RTT 更适合这一版估算。该假设不是最终网络时钟同步方案，"
        "但足以支撑 MVP 级相对进度测量。"
    )

    doc.add_heading("3.3 防止错误样本污染结论", level=2)
    add_bullet(doc, "超过 2500 ms 未更新的 REPORT 被视为陈旧样本，不参与当前成对比较。")
    add_bullet(doc, "两端 PlaybackState 不一致时跳过比较，因为 Playing 与 Paused 的位置推进规律不同。")
    add_bullet(doc, "client_a/client_b 始终按 id 排序，pair_diff_ms 的正负含义不会随触发方变化。")
    add_bullet(doc, "数据锁 mutex_ 与日志锁 logMutex_ 分离；计算后再输出，避免长时间占用共享状态锁。")
    add_bullet(doc, "本阶段只记录日志，不执行 seek 或倍速调整，避免未经验证的指标干扰真实播放。")

    doc.add_heading("3.4 核心修改", level=2)
    add_changed_files_table(doc)

    doc.add_heading("3.5 新日志怎么读", level=2)
    add_code_block(
        doc,
        [
            "[metric] type=pair_progress trigger_client=2 client_a=1 client_b=2",
            "state=Playing projected_a_ms=10406 projected_b_ms=10200",
            "pair_diff_ms=206 pair_abs_diff_ms=206",
            "report_age_a_ms=406 report_age_b_ms=0",
            "one_way_a_ms=0 one_way_b_ms=0 quality=good",
        ],
    )
    add_bullet(doc, "trigger_client：哪一个 client 的新 REPORT 触发了本次比较。")
    add_bullet(doc, "projected_a_ms / projected_b_ms：已经归一化到同一 server 时刻的位置。")
    add_bullet(doc, "pair_diff_ms > 0：client_a 领先；< 0：client_a 落后。")
    add_bullet(doc, "report_age_*：比较时该样本已经存在多久，用于判断样本新鲜度。")
    add_bullet(doc, "quality：当前沿用 good <= 250 ms、watch <= 1000 ms、drift > 1000 ms。")

    doc.add_heading("3.6 本地验证证据", level=2)
    add_numbered(doc, "Windows x64 Debug with libVLC 完整构建通过，SyncCinema.exe 与 SyncCinemaServer.exe 均成功链接。")
    add_numbered(doc, "WSL Ubuntu 22.04 Release server-only 构建通过，验证新增代码没有破坏 Linux 部署目标。")
    add_numbered(doc, "本地启动 server，并用两个独立 TCP client 发送错开 200 ms 的 REPORT。")
    add_numbered(doc, "第一次刻意构造约 206 ms 差距；下一轮经时间投影后得到 -3 ms，符合输入时序。")
    add_callout(
        doc,
        "验收结果",
        "新增指标能够回答“两台 client 当前相差多少毫秒”，且没有修改协议、控制链路或播放器行为。",
    )

    doc.add_heading("3.7 公网真实播放器验证", level=2)
    add_body(
        doc,
        "提交 2c1bf4f 部署到 Linux 云服务器后，使用两份真实 Windows libVLC client 播放同一个 HTTP 视频。"
        "这次验证不是只看“命令有没有广播”，而是同时观察 RTT、client-to-Room 和 pair_progress 三组指标。"
    )
    add_numbered(doc, "两台 client 到云服务器的 RTT 稳定在约 23-28 ms，说明控制链路本身没有明显拥塞。")
    add_numbered(doc, "暂停后两端真实播放器位置分别约为 74651 ms 与 74653 ms，pair_diff 约为 -2 ms。")
    add_numbered(doc, "同一时刻 client-to-Room 仍可达到约 -939 ms，证明 Room 理论时钟的落后量不能直接代表观众间偏差。")
    add_numbered(doc, "播放过程中 pair_diff 常呈现约 250 ms 的台阶，暴露出 libVLC 时间读数粒度、异步播放和每秒上报采样的共同影响。")
    add_callout(
        doc,
        "数据结论",
        "端到端同步观感已经较好，但单条 pair_diff 仍会被采样相位和播放器读数抖动影响。"
        "因此下一阶段不应立即按单个样本 seek，而应先建立控制周期和稳健统计窗口。",
    )

    doc.add_heading("3.8 仍未解决的边界", level=2)
    add_bullet(doc, "新加入房间的 client 还不会自动获取当前房间状态和进度，可能出现数十秒级真实偏差。")
    add_bullet(doc, "PLAY/SEEK 后历史统计尚未按播放 epoch 重置，旧异常值会污染长期平均值。")
    add_bullet(doc, "libVLC 的 play() 是异步请求；“API 返回”不等于首帧已经播放。")
    add_bullet(doc, "min_rtt / 2 假设上下行大致对称；移动网络或拥塞场景下可能不成立。")
    add_bullet(doc, "当前尚未实现自动纠偏、校正冷却、迟滞区间和倍速微调。")

    doc.add_heading("4. 本地阶段：控制 epoch 与稳健偏差窗口", level=1)
    doc.add_heading("4.1 为什么必须划分控制周期", level=2)
    add_body(
        doc,
        "PLAY、PAUSE、SEEK 会改变播放状态或基准位置。若 SEEK 前后的样本继续累计在同一组平均值中，"
        "旧周期的最大值、平均值和最近窗口都会污染新周期，后续自动纠偏就可能依据过期事实做决定。"
        "因此每条有效控制命令都会开启新的 playback epoch。"
    )
    add_code_block(
        doc,
        [
            "PLAY / PAUSE / SEEK accepted",
            "  -> Room applies command and broadcasts it",
            "  -> metrics.beginControlEpoch(command)",
            "  -> epoch += 1",
            "  -> clear progress and pair-window statistics",
            "  -> keep RTT statistics",
        ],
    )
    add_bullet(doc, "RTT 描述网络链路，不会因为一次 seek 自动失效，因此跨 epoch 保留。")
    add_bullet(doc, "进度偏差描述某次控制后的播放结果，必须在新 epoch 中重新采样。")
    add_bullet(doc, "控制命令与 REPORT 共用 controlCommandMutex，避免读到“新 Room 状态 + 旧 epoch”这样的撕裂快照。")

    doc.add_heading("4.2 为什么先等待 2 秒", level=2)
    add_body(
        doc,
        "libVLC 的控制 API 是异步的。命令刚返回时，播放器可能仍在缓冲、解码或寻找关键帧。"
        "新 epoch 的前 2 秒被标记为 settling：日志仍然保留，便于诊断控制响应，"
        "但这些过渡样本不进入长期偏差统计和 pair 窗口。"
    )
    add_callout(
        doc,
        "设计取舍",
        "2 秒是当前 MVP 的可配置经验值，而不是普适常量。后续应根据本地文件、HTTP 视频、"
        "不同编码格式和真实设备数据调整，或改为由播放器事件驱动稳定期结束。",
    )

    doc.add_heading("4.3 为什么选择中位数和 P95", level=2)
    add_body(
        doc,
        "稳定期结束后，每对 client 保存最近 12 个相对偏差。至少积累 6 个样本后窗口才标记为 ready。"
        "单次 250/500 ms 台阶可能只是采样相位抖动，中位数对这种离群值不敏感；"
        "P95 则保留尾部风险，用来观察绝大多数样本之外是否仍有明显卡顿。"
    )
    add_code_block(
        doc,
        [
            "window_size = 12",
            "window_ready = samples >= 6",
            "stable_skew = median(pair_diff_ms)",
            "tail_risk = P95(abs(pair_diff_ms))",
            "severe_streak += 1 only when abs(diff) > 750 ms",
        ],
    )
    add_bullet(doc, "median_diff_ms 保留正负方向：正数表示 client_a 领先，负数表示落后。")
    add_bullet(doc, "median_abs_diff_ms 表示典型偏差大小，可作为未来纠偏阈值的主要输入。")
    add_bullet(doc, "p95_abs_diff_ms 用于识别偶发严重抖动，不能单独决定立即 seek。")
    add_bullet(doc, "consecutive_severe 用于要求异常连续出现，防止一个坏样本触发控制振荡。")

    doc.add_heading("4.4 核心修改", level=2)
    add_epoch_changed_files_table(doc)

    doc.add_heading("4.5 日志字段怎么读", level=2)
    add_code_block(
        doc,
        [
            "[metric] type=control_epoch epoch=2 command=SEEK settle_ms=2000",
            "[metric] type=pair_progress client_a=1 client_b=2 epoch=2",
            "settling=0 window_samples=6 window_ready=1",
            "median_diff_ms=61 median_abs_diff_ms=61 p95_abs_diff_ms=87",
            "consecutive_severe=0 quality=good",
        ],
    )
    add_bullet(doc, "epoch：当前样本属于哪一次控制周期，便于过滤旧数据和复盘命令影响。")
    add_bullet(doc, "settling=1：仍处于播放器过渡期，只观察、不累计、不纠偏。")
    add_bullet(doc, "window_ready=1：样本数量达到最小要求，稳健统计才可供控制算法使用。")
    add_bullet(doc, "quality：窗口未 ready 时仅供观察；ready 后以 median_abs_diff_ms 判定。")

    doc.add_heading("4.6 本地回归结果", level=2)
    add_numbered(doc, "Windows x64 libVLC 完整目标再次构建通过。")
    add_numbered(doc, "WSL Ubuntu 22.04 的 SyncCinemaServer 目标再次构建通过。")
    add_numbered(doc, "两个脚本 TCP client 并发发送 PLAY 和 REPORT，epoch 1 的前 2 秒样本正确标记为 settling。")
    add_numbered(doc, "稳定期后窗口从 1 增长到 6，window_ready 变为 1，中位数约 60-61 ms、P95 约 75-87 ms。")
    add_numbered(doc, "发送 SEEK 后进入 epoch 2，旧窗口被清空并重新进入 settling，验证跨控制周期污染已被阻断。")
    add_callout(
        doc,
        "当前状态",
        "本阶段已作为 4944cf8 发布到 GitHub。它留下的晚加入确定性缺陷，"
        "已经在后续的 92300ce 快照提交中解决。",
    )

    doc.add_heading("4.7 该提交留下的边界", level=2)
    add_bullet(doc, "窗口指标目前只用于日志分析，尚未自动下发校正命令。")
    add_bullet(doc, "4944cf8 尚未处理晚加入 client；该确定性缺陷由下一章的本地候选版本解决。")
    add_bullet(doc, "当前 P95 使用最近 12 个样本的 nearest-rank 算法，样本较少时粒度有限。")
    add_bullet(doc, "暂停和播放共用 2 秒稳定期，后续可按命令类型和播放器事件细化。")

    doc.add_heading("5. 已发布阶段：晚加入快照与房间版本", level=1)
    doc.add_heading("5.1 问题与正确性目标", level=2)
    add_body(
        doc,
        "旧版本 client 连接成功后总是从 Stopped、0 秒开始。若房间已经播放到 80 秒，"
        "它只有等别人再次发送 SEEK 才能追上；这不是网络抖动，而是协议缺少初始化状态。"
        "本阶段先解决这个确定性缺陷，再讨论自动纠偏。",
    )
    add_callout(
        doc,
        "正确性目标",
        "新 client 必须得到同一时刻的房间状态、毫秒位置和控制版本；"
        "在它加入广播列表与完成初始化之间，不能静默漏掉任何 PLAY、PAUSE 或 SEEK。",
    )

    doc.add_heading("5.2 SNAPSHOT 协议与唯一 epoch", level=2)
    add_code_block(
        doc,
        [
            "SNAPSHOT <control_epoch> <Playing|Paused|Stopped> <position_ms>",
            "",
            "example:",
            "SNAPSHOT 7 Playing 123456",
        ],
    )
    add_body(
        doc,
        "Room 成为 controlEpoch 的唯一生产者：每接受一条 PLAY、PAUSE 或 SEEK，epoch 递增一次。"
        "Room::getSnapshot() 在同一把 mutex 下同时读取推算后的 SyncState 和 controlEpoch，"
        "避免把新位置与旧版本拼成一份逻辑上不存在的快照。Metrics 只消费 Room 给出的 epoch，"
        "不再维护第二套独立计数器。",
    )
    add_bullet(doc, "SNAPSHOT 是 server 到 client 的单向初始化消息，client 伪造快照不会改变 Room。")
    add_bullet(doc, "解析器拒绝负 epoch、负位置、非法状态、缺失字段、额外字段和无法安全转换的超大位置。")
    add_bullet(doc, "PLAY/PAUSE/SEEK 的判定集中在 isPlaybackControlMessage()，server、Room 和 metrics 共用一套规则。")

    doc.add_heading("5.3 如何保证加入顺序", level=2)
    add_code_block(
        doc,
        [
            "accept client",
            "  -> lock controlCommandMutex",
            "  -> Room::addClient",
            "  -> Room::getSnapshot",
            "  -> send SNAPSHOT",
            "  -> unlock",
            "  -> start client receive thread",
        ],
    )
    add_body(
        doc,
        "注册、读取快照和发送快照与普通控制命令使用同一把 controlCommandMutex。"
        "因此其他 client 的控制命令只能完整地发生在该流程之前或之后，不能插在中间。"
        "heartbeat 线程仍可能先发出 PING，所以 client 的握手状态机允许在 SNAPSHOT 前响应 PING；"
        "同一次 recv 多读到的后续字节会保留并交给广播线程，TCP 顺序不会被丢弃。",
    )
    add_bullet(doc, "client 最多等待 5 秒；连接旧 server 或协议异常时明确失败，不会永久卡在 recv。")
    add_bullet(doc, "SNAPSHOT 发送失败时，server 会移除连接并关闭 socket，不启动半初始化 client 线程。")
    add_bullet(doc, "快照发送完成后到达的新控制消息会按同一 TCP 连接顺序排在快照后面。")

    doc.add_heading("5.4 真实播放器如何应用快照", level=2)
    add_body(
        doc,
        "openMedia() 成功只代表媒体对象创建完成，网络媒体和真实解码器此时未必可以 seek。"
        "对于非 Stopped+0 的快照，client 先启动播放器，最多等待 15 秒直到 isSeekable()，"
        "再调用 seekMilliseconds()。若快照状态是 Paused 或 Stopped，则定位后暂停；"
        "若是 Playing，则补上从收到快照到实际 seek 的本地经过时间。",
    )
    add_callout(
        doc,
        "精度修正",
        "最初实现把 position_ms 四舍五入成整秒后调用 seek(int)，天然引入最多约 500 ms 误差。"
        "复核时扩展 PlayerController::seekMilliseconds()，libVLC 直接使用原生毫秒时间单位，"
        "MockPlayer 同步保留毫秒状态。",
    )
    add_bullet(doc, "Stopped+0 不启动解码器，避免一个空房间的新连接无意义地弹出播放器窗口。")
    add_bullet(doc, "播放器准备期间 playerMutex 保证同一个 PlayerController 不被其他线程并发控制。")
    add_bullet(doc, "成功后输出 initial_sync 指标，包含 epoch、快照位置、目标位置和应用耗时。")

    doc.add_heading("5.5 核心修改", level=2)
    add_snapshot_changed_files_table(doc)

    doc.add_heading("5.6 验证证据", level=2)
    add_numbered(doc, "Windows x64 Debug with libVLC 与 MockPlayer 两个配置均完整构建通过。")
    add_numbered(doc, "Windows CTest 在两个配置下均为 1/1 通过；WSL Ubuntu 22.04 server 与核心测试构建通过，CTest 1/1 通过。")
    add_numbered(doc, "TCP 晚加入回归：A 发送 SEEK 42 + PLAY 后，B 收到 SNAPSHOT 2 Playing 42359；B 暂停后，C 收到 SNAPSHOT 3 Paused 42360。")
    add_numbered(doc, "MockPlayer 晚加入运行得到 epoch=2、Playing、47086 ms，并正确执行 play 与毫秒 seek。")
    add_numbered(doc, "真实 libVLC 本地文件测试收到 epoch=2、Playing、79863 ms，播放器位置与快照目标均为 79863 ms。")
    add_callout(
        doc,
        "构建诊断",
        "验证时发现目录名虽然是 x64-vlc-debug，旧 CMakeCache 中 USE_LIBVLC 实际为 OFF。"
        "重新执行 cmake --preset x64-vlc-debug 后确认开关为 ON，再完成真实 libVLC 测试。"
        "这说明构建目录名称不是事实来源，缓存变量和运行日志才是。",
    )

    doc.add_heading("5.7 当前局限", level=2)
    add_bullet(doc, "新 client 在首份快照到达前还没有 RTT 样本，初始单向网络传播时间尚未补偿。")
    add_bullet(doc, "该协议升级要求 client/server 同步更新；新 client 连接旧 server 会在 5 秒后主动失败。")
    add_bullet(doc, "普通广播命令尚未携带 epoch，断线重连、去重和过期命令过滤还需要版本化控制信封。")
    add_bullet(doc, "libVLC 冷启动在一次真实测试中约为 16.7 秒，这是产品启动体验问题，不是本阶段快照算法的误差。")
    add_bullet(doc, "本阶段只完成状态初始化，不根据 pair 窗口自动控制播放器。")

    doc.add_heading("6. 已发布阶段：只读校正决策引擎", level=1)
    doc.add_heading("6.1 为什么不直接调用 seek", level=2)
    add_body(
        doc,
        "到 92300ce 为止，系统已经能测量 RTT、把两端上报投影到同一 server 时刻，"
        "并在每个控制 epoch 内形成中位数、P95 和连续严重偏差窗口。"
        "但“测到偏差”并不等于“应该立刻控制播放器”：样本可能仍在稳定期、窗口不足、"
        "方向反复变化，或者刚做过一次校正。若在这些情况下直接 seek，测量噪声会变成用户可见的画面跳动。",
    )
    add_callout(
        doc,
        "本阶段目标",
        "新增一个只读策略层，把窗口证据转换成 hold / would_seek_forward 建议；"
        "它输出原因、目标端和建议前进量，但绝不调用播放器 API。先验证决策质量，再打开闭环控制。",
    )

    doc.add_heading("6.2 测量、决策与执行三层分离", level=2)
    add_code_block(
        doc,
        [
            "SyncMetricsCollector",
            "  -> normalized pair window / RTT / epoch evidence",
            "  -> SyncCorrectionPolicy (pure decision)",
            "  -> correction_advice mode=read_only",
            "",
            "future SyncCorrectionController",
            "  -> validate client/session/epoch again",
            "  -> send correction command",
            "  -> client applies seek and reports result",
        ],
    )
    add_body(
        doc,
        "SyncCorrectionPolicy 不持有共享状态，也不访问 socket、Room 或 PlayerController。"
        "同一组输入总会得到同一决策，因此可以用快速单元测试覆盖边界。"
        "SyncMetricsCollector 仍负责收集和归一化事实；未来控制器才负责命令发送、失败处理和执行反馈。",
    )

    doc.add_heading("6.3 安全门按顺序放行", level=2)
    add_numbered(doc, "稳定期：控制后 2 秒内只观察，返回 settling。")
    add_numbered(doc, "窗口完整性：至少 6 个有效 pair 样本，未满足时返回 window_not_ready。")
    add_numbered(doc, "播放状态：Stopped 不纠偏；Playing 还要求两端均有 RTT，Paused 可直接比较静止位置。")
    add_numbered(doc, "偏差分级：中位绝对偏差不超过 250 ms 时保持；250-750 ms 只观察，不做硬 seek。")
    add_numbered(doc, "方向稳定性：至少 75% 窗口样本必须同向，防止 A/B 领先关系不断翻转。")
    add_numbered(doc, "持续性：至少 3 个连续样本达到 750 ms，排除单点异常。")
    add_numbered(doc, "冷却：一次 would_seek 后模拟 5 秒冷却，验证未来控制节奏不会连续触发。")
    add_body(
        doc,
        "策略以 median_diff_ms 的符号确定领先方向，以 median_abs_diff_ms 判断典型偏差大小。"
        "P95 当前只记录尾部风险，不参与第一版动作门槛；这样可以先积累真实数据，"
        "再决定是否需要把尾部抖动纳入更复杂的策略。",
    )

    doc.add_heading("6.4 为什么只让落后端向前追", level=2)
    add_body(
        doc,
        "pair_diff_ms 定义为 client_a - client_b。若结果为正，说明 B 落后，建议目标是 B；"
        "若为负，说明 A 落后，建议目标是 A。第一版不把领先端向后拉，"
        "因为后退会让用户再次看到已经看过的画面，通常比落后端向前跳更明显。"
        "suggested_forward_ms 使用带符号中位差的绝对值，而不是某个瞬时样本。",
    )
    add_code_block(
        doc,
        [
            "[metric] type=correction_advice mode=read_only",
            "client_a=1 client_b=2 epoch=2 state=Paused",
            "action=would_seek_forward reason=persistent_skew",
            "target_client=2 reference_client=1 suggested_forward_ms=1000",
            "median_diff_ms=1000 direction_agreement_pct=100",
            "consecutive_severe=6 window_samples=6",
        ],
    )

    doc.add_heading("6.5 核心修改", level=2)
    add_correction_changed_files_table(doc)

    doc.add_heading("6.6 验证证据", level=2)
    add_numbered(doc, "Windows x64 Debug、x64 Debug with libVLC 均完整构建，两个配置的 CTest 均为 1/1 通过。")
    add_numbered(doc, "WSL Ubuntu 22.04 的 SyncCinemaServer 与 SyncCinemaCoreTests 构建通过，CTest 1/1 通过。")
    add_numbered(doc, "策略单元测试覆盖所有安全门，以及正负偏差下目标 client、参考 client 和建议前进量。")
    add_numbered(doc, "双脚本 client 集成验证得到 window_not_ready -> persistent_skew -> cooldown 的确定序列。")
    add_numbered(doc, "1 秒稳定偏差场景中，策略在第 6 个有效样本建议 client 2 向前 1000 ms；下一样本进入约 4.9 秒冷却。")
    add_callout(
        doc,
        "关键保证",
        "本地集成测试只验证建议日志，代码路径没有 PlayerController 依赖，也没有执行 seek。"
        "因此这一阶段即使阈值仍需公网调参，也不会改变任何用户播放行为。",
    )

    doc.add_heading("6.7 已知局限与上线前门槛", level=2)
    add_bullet(doc, "当前 5 秒冷却是只读模拟：would_seek 被当作“假设已执行”，用于验证建议节奏。")
    add_bullet(doc, "多于两个 client 时会产生多组 pair 建议；未来执行器需要基于房间共识选目标，不能逐对盲目 seek。")
    add_bullet(doc, "250/750 ms、75% 和 3 个连续样本是保守初值，需要两台真实异地设备的数据校准。")
    add_bullet(doc, "P95 目前只作诊断；是否参与硬校正必须根据误触发率和用户体验决定。")
    add_bullet(doc, "普通控制广播仍未携带命令 id/epoch，真正闭环前需要再次校验连接身份、epoch 和目标状态。")
    add_bullet(doc, "下一步应先部署只读版本采集公网建议日志，再实现 server 到目标 client 的专用校正命令与执行回执。")

    doc.add_heading("7. 本地阶段：Qt 桌面客户端 MVP", level=1)
    doc.add_heading("7.1 为什么先做 UI", level=2)
    add_body(
        doc,
        "同步算法已经具备 RTT、同刻投影、稳健窗口和只读校正建议，但真实异地测试仍依赖对方熟悉命令行。"
        "本阶段先把现有核心能力包装成普通用户可以直接操作的桌面播放器，让第二台电脑能够作为真实测试端参与，"
        "同时坚持不在 UI 文件里复制 TCP 和协议逻辑。",
    )
    add_callout(
        doc,
        "产品目标",
        "用户只需填写同一媒体 URL 与 server 地址，点击连接后即可播放、暂停、拖动进度、调节音量和全屏；"
        "server 仍只协调状态，不传输播放器控制之外的 TCP 视频数据。",
    )

    doc.add_heading("7.2 先抽会话层，再接界面", level=2)
    add_code_block(
        doc,
        [
            "QtClientWindow (widgets / visual state)",
            "  -> QtClientController (Qt <-> C++ adapter)",
            "  -> SyncClientSession (TCP / protocol / worker threads)",
            "  -> PlayerController",
            "       -> LibVlcPlayer",
            "",
            "CLI Client.cpp",
            "  -> same SyncClientSession",
        ],
    )
    add_body(
        doc,
        "旧 Client.cpp 同时承担命令解析、Winsock 生命周期、消息边界、快照握手、广播线程、"
        "心跳回复、进度上报和播放器互斥。若 Qt 再复制一遍，任何协议修复都必须改两处。"
        "因此新增 SyncClientSession，把一个 client 连接的完整生命周期封装为纯 C++ 对象；"
        "CLI 只负责读取 std::cin，Qt 只负责控件和信号，两者通过 callbacks 消费相同事件。",
    )

    doc.add_heading("7.3 UI 线程为什么不能做冷启动和网络握手", level=2)
    add_body(
        doc,
        "libVLC 第一次初始化曾实测需要十几秒，TCP connect 和初始 SNAPSHOT 也可能等待超时。"
        "若这些操作直接发生在 Qt 主线程，窗口会无法拖动、重绘或响应关闭。QtClientController 因此使用"
        " startupThread 创建播放器、打开媒体和建立会话；网络回调通过 QMetaObject::invokeMethod"
        " 的 QueuedConnection 回到 UI 线程，再更新控件。后台线程从不直接操作 QWidget。",
    )
    add_bullet(doc, "resourcesMutex 保护 controller 持有的 player/session 指针和断开时的所有权转移。")
    add_bullet(doc, "SyncClientSession 内部 playerMutex 保护同一个播放器，sendMutex 保证多线程发送不会交叉 TCP 字节。")
    add_bullet(doc, "关闭窗口时先取消启动、停止会话并 join 线程，再销毁被 session 引用的 player。")

    doc.add_heading("7.4 播放器嵌入与常用控件", level=2)
    add_body(
        doc,
        "Qt 的视频区域设置为原生窗口，winId() 转成 Windows HWND 后交给"
        " libvlc_media_player_set_hwnd。画面由 libVLC 渲染，Qt 负责其外部布局。"
        "PlayerController 新增 setVideoOutputWindow、getDurationMilliseconds、setVolume 和 getVolume，"
        "因此 Qt 层仍然不接触 libVLC 头文件之外的实现细节。",
    )
    add_bullet(doc, "顶部保留媒体文件/URL、文件选择、server 地址、连接按钮和状态指示。")
    add_bullet(doc, "视频区域全宽占据主体，不使用营销式卡片或无关装饰。")
    add_bullet(doc, "底部采用播放器熟悉的播放/暂停、时间轴、时间、音量和全屏图标；双击画面也可全屏，Esc 退出。")
    add_bullet(doc, "QSettings 保存媒体地址、server 地址和音量，降低重复测试成本。")

    doc.add_heading("7.5 构建与发布边界", level=2)
    add_body(
        doc,
        "BUILD_QT_CLIENT 默认 OFF，保证 Linux 云 server 和未安装 Qt 的开发环境不受影响。"
        "x64-qt-vlc-release preset 才启用桌面目标；CMake 优先查找 Qt 6，找不到时兼容 Qt 5。"
        "构建后 windeployqt 自动复制 Qt DLL 与 platform plugin，公共 helper 同时复制 libVLC DLL 和 plugins。"
        "给测试用户时必须发送整个输出目录，而不是单独发送 exe。",
    )
    add_callout(
        doc,
        "环境兼容记录",
        "本机 Anaconda Qt 5.15 仍使用已被 VS 2026 删除的 stdext 检查迭代器。"
        "QtMsvcCompatibility.h 只在 Qt 5 + MSVC 下把相关宏退化为普通指针；"
        "这是当前开发环境适配，不是长期替代官方 Qt 6 SDK 的方案。",
    )

    doc.add_heading("7.6 核心修改", level=2)
    add_qt_changed_files_table(doc)

    doc.add_heading("7.7 当前验证证据", level=2)
    add_numbered(doc, "SyncCinemaQt、libVLC CLI、SyncCinemaServer 和 SyncCinemaCoreTests 在 Windows Release 配置共同构建通过。")
    add_numbered(doc, "无 libVLC 的 Windows x64 Debug MockPlayer 配置构建通过，CTest 1/1 通过。")
    add_numbered(doc, "WSL Ubuntu 22.04 的 Linux server 和核心测试构建通过，CTest 1/1 通过。")
    add_numbered(doc, "CLI 通过新 SyncClientSession 完成真实 TCP 连接、SNAPSHOT 0 Stopped 0 握手、状态查询和正常退出。")
    add_numbered(doc, "Qt 界面在 1180x760 与最小 820x560 下完成真实渲染截图检查，无控件重叠、文字截断或空白窗口。")

    doc.add_heading("7.8 当前局限", level=2)
    add_bullet(doc, "当前尚未在第二台真实电脑上完成 Qt 客户端的完整播放、暂停和 seek 验收。")
    add_bullet(doc, "server 端口仍固定为 9000，界面尚无房间码、用户身份和在线成员列表。")
    add_bullet(doc, "聊天、弹幕、安装器、自动更新、日志文件和崩溃收集均不属于第一版。")
    add_bullet(doc, "只读校正建议仍未闭环执行；UI 不会自动 seek 用户播放器。")
    add_bullet(doc, "公开发布前应换用正式 Qt SDK，并建立可复现的 Release 打包与签名流程。")

    doc.add_heading("8. 面试表达模板", level=1)
    add_body(
        doc,
        "第一段可以这样说明指标演进：最初我用服务器 Room 的理论时钟评估每个客户端，"
        "但公网测试发现两端都可能因 libVLC 缓冲共同落后 Room 几百毫秒，"
        "这个指标不能直接代表观众之间不同步。于是我保留 client-to-Room 指标用于诊断，"
        "另外缓存每个 client 的最近上报，结合 RTT/2 和 server steady_clock，"
        "把异步到达的进度投影到同一时刻后计算 pair_diff。"
        "同时过滤陈旧样本和状态不一致样本，并坚持先观测后控制，"
        "为后续自动纠偏建立可信数据基础。"
    )
    add_body(
        doc,
        "第二段可以这样说明稳健性：公网实测中，两台客户端暂停后只差约 2 ms，"
        "但播放时单条日志会出现约 250 ms 的量化台阶。我没有直接按单个样本 seek，"
        "而是为每次 PLAY、PAUSE、SEEK 建立独立 epoch，清理旧进度统计但保留 RTT；"
        "控制后设置 2 秒稳定期，再用最近 12 个样本的中位数评估典型偏差、P95 观察尾部风险，"
        "同时统计连续严重异常。这样未来的校正策略会基于趋势，而不是被播放器异步行为或一次网络抖动误触发。"
    )
    add_body(
        doc,
        "第三段可以这样说明晚加入：自动纠偏前我先处理了一个确定性更强的问题，"
        "即新客户端加入正在播放的房间时会停在 0 秒。我让 Room 统一维护 controlEpoch，"
        "新增 SNAPSHOT 消息携带状态、毫秒位置和版本；server 用控制顺序锁把注册、取快照和发送变成一个有序步骤。"
        "client 侧实现带超时的握手，能在快照前响应心跳并保留多读字节；真实播放器等待 seekable 后按毫秒定位。"
        "同时用 Windows/Linux 核心测试、脚本双客户端和真实 libVLC 三层验证，确保协议、并发顺序和播放器行为都成立。"
    )
    add_body(
        doc,
        "第四段可以这样说明控制安全性：有了相对偏差后我没有直接 seek，而是把测量、决策和执行拆成三层。"
        "纯策略层按稳定期、窗口、状态、RTT、阈值、方向一致率、连续异常和冷却逐层放行，"
        "只输出 would_seek_forward 及可解释 reason。它只让落后端向前追赶，并通过跨平台单元测试和双客户端集成测试"
        "验证 window_not_ready、persistent_skew、cooldown 的状态序列。当前仍是 read_only，"
        "先用公网数据验证误触发率，再接入真正的命令执行和回执。"
    )

    add_body(
        doc,
        "第五段可以这样说明客户端工程化：为了让另一台电脑由普通用户直接参与测试，"
        "我没有在 Qt 中重写一套网络代码，而是先把旧 Client.cpp 拆成纯 C++ 的 SyncClientSession。"
        "CLI 和 Qt 共用快照握手、心跳、广播与进度上报；QtClientController 在后台完成 libVLC 冷启动和连接，"
        "再用 queued invoke 回到 UI 线程。视频通过 HWND 嵌入，CMake 把 Qt 目标设为可选并自动部署 Qt/VLC runtime，"
        "因此 Linux server 构建路径保持不变。",
    )

    doc.add_heading("9. 后续学习重点", level=1)
    add_numbered(doc, "SyncClientSession：如何把 UI/CLI 与网络、协议、播放器并发解耦。")
    add_numbered(doc, "QtClientController：为什么慢操作放到 worker thread，QWidget 更新必须回到 UI 线程。")
    add_numbered(doc, "QObject signal、QMetaObject::invokeMethod 与 QueuedConnection 的线程边界。")
    add_numbered(doc, "libvlc_media_player_set_hwnd：第三方播放器如何嵌入 Qt 原生窗口。")
    add_numbered(doc, "CMake 的可选目标、find_package、AUTOMOC 和 POST_BUILD runtime 部署。")
    add_numbered(doc, "evaluateSyncCorrection：如何用提前返回把安全门写成可读、可测试的纯决策函数。")
    add_numbered(doc, "calculateDirectionAgreementPercent：为什么偏差大小之外还必须验证方向稳定。")
    add_numbered(doc, "PairWindowStats 的建议去重与模拟冷却：状态为什么属于 pair+epoch，而不是某个 client。")
    add_numbered(doc, "策略单元测试：如何用输入夹具覆盖安全门和正负方向，而不启动网络与播放器。")
    add_numbered(doc, "SyncServer 的 accept 初始化临界区：为什么 addClient、getSnapshot、send 必须与控制命令共用顺序锁。")
    add_numbered(doc, "SyncClientSession::receiveInitialSnapshot：select 超时、PING 兼容和 TCP 多读缓冲如何组成握手状态机。")
    add_numbered(doc, "SyncClientSession::applyInitialSnapshot：播放器准备、毫秒 seek 和 Playing 状态经过时间补偿。")
    add_numbered(doc, "Room::getSnapshot：如何在同一把锁下形成 state+epoch 的一致快照。")
    add_numbered(doc, "SyncMetricsCollector::recordProgressReport：共享状态、快照、时间投影和日志输出如何分层。")
    add_numbered(doc, "projectPositionToServerTime：为什么 Playing 才能按经过时间推进。")
    add_numbered(doc, "SyncMetricsCollector::beginControlEpoch：为什么 RTT 保留，而播放偏差必须按控制周期重置。")
    add_numbered(doc, "medianOf / percentile95OfAbsoluteDiffs：稳健统计如何减少离群样本误导。")
    add_numbered(doc, "SyncServer::processLine：controlCommandMutex 如何统一 Room、广播、epoch 和 REPORT 的并发顺序。")
    add_numbered(doc, "Room::getEstimatedStateLocked：服务器如何用基准位置与 steady_clock 维护权威状态。")
    add_numbered(doc, "SyncMetricsCollector::recordPongReceived：为什么 RTT 可以只用 server 自己的时钟计算。")
    add_numbered(doc, "SyncClientSession 的 playerMutex/sendMutex：多个线程如何安全共享播放器和同一个 TCP 字节流。")

    doc.add_heading("10. 后续工程路线", level=1)
    add_numbered(doc, "在合理提交间隔后提交 Qt 桌面客户端候选，并打包完整 runtime 目录交给第二台电脑做真实用户验收。")
    add_numbered(doc, "用两台真实设备采集 read_only 建议日志，统计建议频率、方向正确率和误触发场景。")
    add_numbered(doc, "设计专用校正协议：携带目标 client、room epoch、目标毫秒位置和命令 id，并要求应用结果回执。")
    add_numbered(doc, "加入最小闭环硬 seek：server 只向落后端发送，client 再次校验 epoch，执行后上报实际位置；失败时不连续重试。")
    add_numbered(doc, "为普通控制广播加入 epoch/命令 id，完善重连后的过期消息过滤和幂等语义。")
    add_numbered(doc, "云端闭环稳定后，再研究 250-750 ms 小偏差的温和校正；暂不贸然引入倍速控制。")
    add_numbered(doc, "进一步同步 server/client 时钟并支持未来执行时间，让两端在约定时刻执行控制，降低广播到达差。")
    add_numbered(doc, "在 Qt 客户端上继续增加房间信息、聊天与弹幕，但不让 UI 直接处理协议字符串。")
    add_numbered(doc, "补齐 systemd 服务、配置文件、日志轮转、自动构建测试和发布包。")

    add_callout(
        doc,
        "下一阶段入口",
        "只读校正策略已作为 2f4f27a 发布。feature/qt-client-ui 已形成可构建、可打包的本地候选，"
        "并保持 CLI 与 Linux server 回归通过。下一步在合理提交间隔后提交该 UI 阶段，"
        "把完整运行目录交给第二台电脑做真实用户测试；随后用异地日志校准阈值，再打开最小闭环。",
    )

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
